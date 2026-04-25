import json
import os
import re
import docx.shared
from docx import Document
from docx.shared import Pt

def limpar_texto_xml(texto):
    """
    Remove caracteres de controle invisíveis que quebram o XML do Word.
    """
    if not isinstance(texto, str):
        return ""
    return re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', texto)

def formatar_documento_comentarios(caminho_json="data/resultado_busca.json", caminho_saida="data/Comentarios.docx"):
    print(f"\nFormatando documento combinando resultados...")
    
    # Descobre a pasta onde os arquivos JSON estão salvos
    pasta_data = os.path.dirname(caminho_json) or "."
    caminho_tec = os.path.join(pasta_data, "resultado_busca_tec.json")
    caminho_qc = os.path.join(pasta_data, "resultado_busca_qc.json")
    
    resultados_tec = []
    resultados_qc = []
    
    # Lê os resultados do TEC, se existirem
    if os.path.exists(caminho_tec):
        with open(caminho_tec, 'r', encoding='utf-8') as f:
            resultados_tec = json.load(f)
            
    # Lê os resultados do QC, se existirem
    if os.path.exists(caminho_qc):
        with open(caminho_qc, 'r', encoding='utf-8') as f:
            resultados_qc = json.load(f)

    # Fallback caso os nomes não sejam os esperados
    if not resultados_tec and not resultados_qc:
        if os.path.exists(caminho_json):
            with open(caminho_json, 'r', encoding='utf-8') as f:
                resultados_tec = json.load(f)
        else:
            print("Nenhum arquivo JSON encontrado para formatar!")
            return False

    # Usar a lista que tiver mais itens como base (garante que vamos iterar sobre todas as questões)
    lista_base = resultados_tec if len(resultados_tec) >= len(resultados_qc) else resultados_qc
    
    if not lista_base:
        print("A lista de resultados está vazia.")
        return False

    doc = Document()
    
    # Configura a fonte padrão
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    # Inicia o bloco de questões
    doc.add_paragraph("<bloco_de_questoes>")

    for idx, item in enumerate(lista_base, 1):
        texto_original = item.get("texto_original", "")
        
        # Busca as correspondências em ambas as listas comparando o enunciado original
        item_tec = next((q for q in resultados_tec if q.get("texto_original") == texto_original), None)
        item_qc = next((q for q in resultados_qc if q.get("texto_original") == texto_original), None)
        
        # Abertura do Item
        doc.add_paragraph(f'<item_questao numero="{idx}">')
        
        # Enunciado
        doc.add_paragraph("<enunciado_questao>")
        doc.add_paragraph(limpar_texto_xml(texto_original))
        doc.add_paragraph("</enunciado_questao>")
        
        # Comentários
        doc.add_paragraph("<comentarios_selecionados>")
        
        # =======================================================
        # COMENTÁRIO 1.1 (TEC)
        # =======================================================
        p_coment1 = doc.add_paragraph()
        run_c1 = p_coment1.add_run(f"## **COMENTÁRIO {idx}.1**")
        run_c1.bold = True
        
        if item_tec:
            status = item_tec.get("status", "")
            comentario_texto = limpar_texto_xml(item_tec.get("comentario", ""))
            
            p_texto = doc.add_paragraph()
            if status == "Não encontrada":
                run_err = p_texto.add_run("Informação: A questão não foi encontrada no TEC concursos.")
                run_err.font.color.rgb = docx.shared.RGBColor(255, 0, 0)
            else:
                p_texto.add_run(comentario_texto)
        else:
            p_texto = doc.add_paragraph()
            run_err = p_texto.add_run("[Busca no TEC não realizada para esta questão]")
            run_err.font.color.rgb = docx.shared.RGBColor(128, 128, 128) # Cinza
            
        # =======================================================
        # COMENTÁRIO 1.2 (QCONCURSOS)
        # =======================================================
        p_coment2 = doc.add_paragraph()
        run_c2 = p_coment2.add_run(f"## **COMENTÁRIO {idx}.2**")
        run_c2.bold = True
        
        if item_qc:
            status = item_qc.get("status", "")
            comentario_texto = limpar_texto_xml(item_qc.get("comentario", ""))
            
            p_texto = doc.add_paragraph()
            if status == "Não encontrada":
                run_err = p_texto.add_run("Informação: A questão não foi encontrada no QConcursos.")
                run_err.font.color.rgb = docx.shared.RGBColor(255, 0, 0)
            else:
                p_texto.add_run(comentario_texto)
        else:
            # Mantém o texto padrão de placeholder caso não tenha busca no QC
            doc.add_paragraph("(((QCONCURSOS)))")
        
        # Fechamento de Comentários e Inserção de Gabarito
        doc.add_paragraph("</comentarios_selecionados>")
        
        doc.add_paragraph("<gabarito>")
        doc.add_paragraph("[Gabarito: ]")
        doc.add_paragraph("</gabarito>")
        
        # Fechamento do Item
        doc.add_paragraph("</item_questao>")

    # Finaliza o bloco de questões
    doc.add_paragraph("</bloco_de_questoes>")

    os.makedirs(os.path.dirname(caminho_saida), exist_ok=True)
    doc.save(caminho_saida)
    print(f"✅ Documento formatado com sucesso: {caminho_saida}")
    return True